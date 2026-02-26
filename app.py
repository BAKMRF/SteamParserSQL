import streamlit as st
import pandas as pd
import requests
import os
import re
import time
from datetime import datetime, timedelta
import plotly.express as px
import plotly.graph_objects as go
import numpy as np
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import json
from pathlib import Path
import psycopg2
from psycopg2.extras import RealDictCursor
from contextlib import contextmanager
from dotenv import load_dotenv

# ==================== КОНФИГУРАЦИЯ ====================

load_dotenv()

# Настройки PostgreSQL
DB_CONFIG = {
    'host': os.getenv('DB_HOST', 'localhost'),
    'port': os.getenv('DB_PORT', '5432'),
    'database': os.getenv('DB_NAME', 'steam_parser'),
    'user': os.getenv('DB_USER', 'steam_user'),
    'password': os.getenv('DB_PASSWORD', 'steam_password')
}

STEAM_API_KEY = os.getenv("STEAM_API_KEY", "")
DEMO_MODE = False  # ВСЕГДА РЕАЛЬНЫЙ РЕЖИМ!

STEAM_ACCOUNTS = [
    'https://steamcommunity.com/profiles/76561199001022272',
    'https://steamcommunity.com/profiles/76561199219594998',
    'https://steamcommunity.com/profiles/76561199384092020',
    'https://steamcommunity.com/profiles/76561198333882340',
    'https://steamcommunity.com/profiles/76561199038225456',
    'https://steamcommunity.com/profiles/76561199082417445',
]

# Папки для хранения данных
BASE_DIR = Path(__file__).parent
DATA_DIR = BASE_DIR / "data"
WORD_REPORTS_DIR = DATA_DIR / "word_reports"

# Создаем директории если их нет
DATA_DIR.mkdir(exist_ok=True)
WORD_REPORTS_DIR.mkdir(exist_ok=True)

# ==================== КЛАСС ДЛЯ РАБОТЫ С БД ====================

class DatabaseManager:
    def __init__(self, config):
        self.config = config
        self._init_db()
    
    @contextmanager
    def get_connection(self):
        """Контекстный менеджер для подключения к БД"""
        conn = None
        try:
            conn = psycopg2.connect(**self.config)
            yield conn
        except Exception as e:
            print(f"❌ Ошибка подключения к БД: {e}")
            raise
        finally:
            if conn:
                conn.close()
    
    @contextmanager
    def get_cursor(self, cursor_factory=RealDictCursor):
        """Контекстный менеджер для курсора"""
        with self.get_connection() as conn:
            cursor = conn.cursor(cursor_factory=cursor_factory)
            try:
                yield cursor
                conn.commit()
            except Exception as e:
                conn.rollback()
                print(f"❌ Ошибка выполнения запроса: {e}")
                raise
            finally:
                cursor.close()
    
    def _init_db(self):
        """Проверка подключения к БД и создание таблиц если их нет"""
        try:
            with self.get_cursor() as cursor:
                # Проверяем существование таблиц
                cursor.execute("""
                    SELECT EXISTS (
                        SELECT FROM information_schema.tables 
                        WHERE table_name = 'parse_sessions'
                    );
                """)
                tables_exist = cursor.fetchone()['exists']
                
                if not tables_exist:
                    print("🔄 Создание таблиц в базе данных...")
                    self._create_tables()
                else:
                    print("✅ Подключение к PostgreSQL установлено")
                    
        except Exception as e:
            print(f"❌ Ошибка подключения к PostgreSQL: {e}")
            print("Убедитесь, что Docker контейнер запущен: docker-compose up -d")
    
    def _create_tables(self):
        """Создает таблицы в базе данных"""
        create_tables_sql = """
        -- Создание enum типов
        DO $$ BEGIN
            CREATE TYPE parse_status AS ENUM ('success', 'failed', 'pending');
        EXCEPTION
            WHEN duplicate_object THEN null;
        END $$;

        -- Таблица для сессий парсинга
        CREATE TABLE IF NOT EXISTS parse_sessions (
            id SERIAL PRIMARY KEY,
            parse_time TIMESTAMP WITH TIME ZONE DEFAULT CURRENT_TIMESTAMP,
            parse_date DATE,
            parse_time_display VARCHAR(50),
            timestamp_str VARCHAR(20),
            total_profiles INTEGER DEFAULT 0,
            successful_profiles INTEGER DEFAULT 0,
            failed_profiles INTEGER DEFAULT 0,
            status VARCHAR(20) DEFAULT 'pending',
            created_at TIMESTAMP WITH TIME ZONE DEFAULT CURRENT_TIMESTAMP,
            updated_at TIMESTAMP WITH TIME ZONE DEFAULT CURRENT_TIMESTAMP
        );

        -- Таблица для профилей
        CREATE TABLE IF NOT EXISTS profiles (
            id SERIAL PRIMARY KEY,
            steam_id VARCHAR(50) UNIQUE NOT NULL,
            nickname VARCHAR(255),
            country VARCHAR(10),
            avatar_url TEXT,
            steam_level INTEGER DEFAULT 0,
            profile_url TEXT,
            first_seen TIMESTAMP WITH TIME ZONE DEFAULT CURRENT_TIMESTAMP,
            last_updated TIMESTAMP WITH TIME ZONE DEFAULT CURRENT_TIMESTAMP
        );

        -- Таблица для данных парсинга
        CREATE TABLE IF NOT EXISTS profile_snapshots (
            id SERIAL PRIMARY KEY,
            session_id INTEGER REFERENCES parse_sessions(id) ON DELETE CASCADE,
            profile_id INTEGER REFERENCES profiles(id) ON DELETE CASCADE,
            steam_level INTEGER DEFAULT 0,
            games_count INTEGER DEFAULT 0,
            library_value DECIMAL(10, 2) DEFAULT 0,
            inventory_value DECIMAL(10, 2) DEFAULT 0,
            total_value DECIMAL(10, 2) DEFAULT 0,
            parsed_at TIMESTAMP WITH TIME ZONE DEFAULT CURRENT_TIMESTAMP,
            status VARCHAR(20) DEFAULT 'success',
            error_message TEXT,
            UNIQUE(session_id, profile_id)
        );

        -- Создание индексов
        CREATE INDEX IF NOT EXISTS idx_parse_sessions_parse_date ON parse_sessions(parse_date);
        CREATE INDEX IF NOT EXISTS idx_parse_sessions_parse_time ON parse_sessions(parse_time);
        CREATE INDEX IF NOT EXISTS idx_profile_snapshots_session_id ON profile_snapshots(session_id);
        CREATE INDEX IF NOT EXISTS idx_profile_snapshots_profile_id ON profile_snapshots(profile_id);
        CREATE INDEX IF NOT EXISTS idx_profiles_steam_id ON profiles(steam_id);
        CREATE INDEX IF NOT EXISTS idx_profiles_last_updated ON profiles(last_updated);

        -- Функция для обновления updated_at
        CREATE OR REPLACE FUNCTION update_updated_at_column()
        RETURNS TRIGGER AS $$
        BEGIN
            NEW.updated_at = CURRENT_TIMESTAMP;
            RETURN NEW;
        END;
        $$ language 'plpgsql';

        -- Триггеры
        DROP TRIGGER IF EXISTS update_parse_sessions_updated_at ON parse_sessions;
        CREATE TRIGGER update_parse_sessions_updated_at 
            BEFORE UPDATE ON parse_sessions 
            FOR EACH ROW 
            EXECUTE FUNCTION update_updated_at_column();

        DROP TRIGGER IF EXISTS update_profiles_updated_at ON profiles;
        CREATE TRIGGER update_profiles_updated_at 
            BEFORE UPDATE ON profiles 
            FOR EACH ROW 
            EXECUTE FUNCTION update_updated_at_column();

        -- Представление для удобной агрегации
        DROP VIEW IF EXISTS session_summary;
        CREATE VIEW session_summary AS
        SELECT 
            ps.id as session_id,
            ps.parse_time,
            ps.parse_date,
            ps.parse_time_display,
            ps.total_profiles,
            ps.successful_profiles,
            ps.failed_profiles,
            ps.status,
            COUNT(DISTINCT p.country) as countries_count,
            COALESCE(SUM(psnap.games_count), 0) as total_games,
            COALESCE(AVG(psnap.steam_level), 0)::NUMERIC(10,2) as avg_level,
            COALESCE(SUM(psnap.library_value), 0) as total_library_value,
            COALESCE(SUM(psnap.inventory_value), 0) as total_inventory_value,
            COALESCE(SUM(psnap.library_value + psnap.inventory_value), 0) as grand_total_value
        FROM parse_sessions ps
        LEFT JOIN profile_snapshots psnap ON ps.id = psnap.session_id
        LEFT JOIN profiles p ON psnap.profile_id = p.id
        GROUP BY ps.id, ps.parse_time, ps.parse_date, ps.parse_time_display, 
                 ps.total_profiles, ps.successful_profiles, ps.failed_profiles, ps.status;
        """
        
        try:
            with self.get_cursor() as cursor:
                cursor.execute(create_tables_sql)
            print("✅ Таблицы успешно созданы")
        except Exception as e:
            print(f"❌ Ошибка при создании таблиц: {e}")
            # Не вызываем исключение, чтобы приложение продолжило работу
    
    def create_parse_session(self, parse_time=None):
        """Создает новую сессию парсинга"""
        if not parse_time:
            parse_time = datetime.now()
        
        parse_time_display = parse_time.strftime("%d.%m.%Y %H:%M:%S")
        timestamp_str = parse_time.strftime("%Y%m%d_%H%M%S")
        
        with self.get_cursor() as cursor:
            cursor.execute("""
                INSERT INTO parse_sessions 
                (parse_time, parse_time_display, timestamp_str, status)
                VALUES (%s, %s, %s, 'pending')
                RETURNING id
            """, (parse_time, parse_time_display, timestamp_str))
            
            result = cursor.fetchone()
            return result['id'] if result else None
    
    def update_session_stats(self, session_id, total_profiles, successful, failed):
        """Обновляет статистику сессии"""
        with self.get_cursor() as cursor:
            cursor.execute("""
                UPDATE parse_sessions 
                SET total_profiles = %s,
                    successful_profiles = %s,
                    failed_profiles = %s,
                    status = CASE 
                        WHEN %s > 0 THEN 'success'
                        ELSE 'failed'
                    END
                WHERE id = %s
            """, (total_profiles, successful, failed, successful, session_id))
    
    def get_or_create_profile(self, steam_id, profile_data):
        """Получает или создает профиль"""
        with self.get_cursor() as cursor:
            # Пытаемся найти существующий профиль
            cursor.execute("""
                SELECT id FROM profiles WHERE steam_id = %s
            """, (steam_id,))
            
            result = cursor.fetchone()
            
            if result:
                profile_id = result['id']
                # Обновляем данные профиля
                cursor.execute("""
                    UPDATE profiles 
                    SET nickname = %s,
                        country = %s,
                        avatar_url = %s,
                        steam_level = %s,
                        profile_url = %s,
                        last_updated = CURRENT_TIMESTAMP
                    WHERE id = %s
                """, (
                    profile_data.get('nickname'),
                    profile_data.get('country'),
                    profile_data.get('avatar'),
                    profile_data.get('steam_level'),
                    profile_data.get('profile_url'),
                    profile_id
                ))
                return profile_id
            else:
                # Создаем новый профиль
                cursor.execute("""
                    INSERT INTO profiles 
                    (steam_id, nickname, country, avatar_url, steam_level, profile_url)
                    VALUES (%s, %s, %s, %s, %s, %s)
                    RETURNING id
                """, (
                    steam_id,
                    profile_data.get('nickname'),
                    profile_data.get('country'),
                    profile_data.get('avatar'),
                    profile_data.get('steam_level'),
                    profile_data.get('profile_url')
                ))
                result = cursor.fetchone()
                return result['id'] if result else None
    
    def save_profile_snapshot(self, session_id, profile_id, profile_data, status='success', error=None):
        """Сохраняет снимок данных профиля"""
        with self.get_cursor() as cursor:
            cursor.execute("""
                INSERT INTO profile_snapshots 
                (session_id, profile_id, steam_level, games_count, 
                 library_value, inventory_value, parsed_at, status, error_message)
                VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s)
                ON CONFLICT (session_id, profile_id) 
                DO UPDATE SET
                    steam_level = EXCLUDED.steam_level,
                    games_count = EXCLUDED.games_count,
                    library_value = EXCLUDED.library_value,
                    inventory_value = EXCLUDED.inventory_value,
                    parsed_at = EXCLUDED.parsed_at,
                    status = EXCLUDED.status,
                    error_message = EXCLUDED.error_message
            """, (
                session_id, profile_id,
                profile_data.get('steam_level', 0),
                profile_data.get('games_count', 0),
                profile_data.get('library_value', 0),
                profile_data.get('inventory_value', 0),
                profile_data.get('parsed_at', datetime.now()),
                status, error
            ))
    
    def get_sessions(self, limit=100):
        """Получает список сессий парсинга"""
        with self.get_cursor() as cursor:
            cursor.execute("""
                SELECT * FROM session_summary 
                ORDER BY parse_time DESC 
                LIMIT %s
            """, (limit,))
            return cursor.fetchall()
    
    def get_session_by_id(self, session_id):
        """Получает данные сессии по ID"""
        with self.get_cursor() as cursor:
            cursor.execute("""
                SELECT * FROM session_summary WHERE session_id = %s
            """, (session_id,))
            return cursor.fetchone()
    
    def get_session_profiles(self, session_id):
        """Получает все профили для конкретной сессии"""
        with self.get_cursor() as cursor:
            cursor.execute("""
                SELECT 
                    p.*,
                    ps.steam_level as snapshot_level,
                    ps.games_count,
                    ps.library_value,
                    ps.inventory_value,
                    ps.total_value,
                    ps.parsed_at,
                    ps.status as snapshot_status,
                    ps.error_message
                FROM profile_snapshots ps
                JOIN profiles p ON ps.profile_id = p.id
                WHERE ps.session_id = %s
                ORDER BY ps.total_value DESC
            """, (session_id,))
            return cursor.fetchall()
    
    def get_profile_history(self, profile_id, limit=50):
        """Получает историю изменений профиля"""
        with self.get_cursor() as cursor:
            cursor.execute("""
                SELECT 
                    s.parse_time,
                    s.parse_time_display,
                    ps.*
                FROM profile_snapshots ps
                JOIN parse_sessions s ON ps.session_id = s.id
                WHERE ps.profile_id = %s
                ORDER BY s.parse_time DESC
                LIMIT %s
            """, (profile_id, limit))
            return cursor.fetchall()
    
    def delete_session(self, session_id):
        """Удаляет сессию и связанные данные"""
        with self.get_cursor() as cursor:
            cursor.execute("DELETE FROM parse_sessions WHERE id = %s", (session_id,))
    
    def get_stats(self):
        """Получает общую статистику по БД"""
        with self.get_cursor() as cursor:
            cursor.execute("""
                SELECT 
                    (SELECT COUNT(*) FROM profiles) as total_profiles,
                    (SELECT COUNT(*) FROM parse_sessions) as total_sessions,
                    (SELECT COUNT(*) FROM profile_snapshots) as total_snapshots,
                    (SELECT MAX(parse_time) FROM parse_sessions) as last_parse,
                    (SELECT SUM(games_count) FROM profile_snapshots) as total_games,
                    (SELECT SUM(total_value) FROM profile_snapshots) as total_value
            """)
            return cursor.fetchone()

# ==================== КЛАСС ПАРСЕРА ====================

class SteamParser:
    def __init__(self, db_manager):
        self.api_key = STEAM_API_KEY
        self.demo_mode = DEMO_MODE
        self.db = db_manager
        
        print(f"🔧 Инициализация парсера:")
        print(f"   Режим: {'ДЕМО' if self.demo_mode else 'РЕАЛЬНЫЙ'}")
        print(f"   API ключ: {'ЕСТЬ' if self.api_key else 'НЕТ'}")
        
    def extract_steam_id(self, input_str: str) -> str:
        """Извлекает SteamID из разных форматов"""
        if re.match(r'^\d{17}$', input_str):
            return input_str
        
        match = re.search(r'steamcommunity\.com/(?:profiles|id)/([a-zA-Z0-9_]+)', input_str)
        if match:
            if not match.group(1).isdigit():
                return self._resolve_vanity_url(match.group(1))
            return match.group(1)
        
        match = re.search(r'steamcommunity\.com/profiles/(\d{17})', input_str)
        if match:
            return match.group(1)
        
        return input_str
    
    def _resolve_vanity_url(self, vanity_name: str) -> str:
        """Преобразует никнейм в SteamID"""
        try:
            url = "https://api.steampowered.com/ISteamUser/ResolveVanityURL/v1/"
            params = {'key': self.api_key, 'vanityurl': vanity_name}
            
            response = requests.get(url, params=params, timeout=10)
            data = response.json()
            
            if data['response']['success'] == 1:
                return data['response']['steamid']
        except:
            pass
        return ""
    
    def get_player_info(self, steam_id: str) -> dict:
        """Получает информацию об игроке через Steam API"""
        if not self.api_key:
            print(f"❌ Нет API ключа! Используем заглушку для {steam_id}")
            return {
                'personaname': f'User_{steam_id[-8:]}',
                'loccountrycode': 'RU',
                'avatarfull': '',
                'profileurl': f'https://steamcommunity.com/profiles/{steam_id}',
                'steamid': steam_id
            }
        
        try:
            print(f"🌐 Запрос реальных данных для {steam_id}")
            url = "https://api.steampowered.com/ISteamUser/GetPlayerSummaries/v2/"
            params = {'key': self.api_key, 'steamids': steam_id}
            
            response = requests.get(url, params=params, timeout=10)
            
            if response.status_code == 200:
                data = response.json()
                
                if data['response']['players']:
                    player = data['response']['players'][0]
                    print(f"✅ Получены реальные данные для {steam_id}: {player.get('personaname', 'Unknown')}")
                    return player
                else:
                    print(f"⚠️  Нет данных об игроке {steam_id}")
            else:
                print(f"❌ Ошибка HTTP {response.status_code} для {steam_id}")
                
        except Exception as e:
            print(f"❌ Ошибка API для {steam_id}: {str(e)}")
        
        return {
            'personaname': f'User_{steam_id[-8:]}',
            'loccountrycode': 'Unknown',
            'avatarfull': '',
            'profileurl': f'https://steamcommunity.com/profiles/{steam_id}',
            'steamid': steam_id
        }
    
    def get_steam_level(self, steam_id: str) -> int:
        """Получает уровень Steam аккаунта"""
        if not self.api_key:
            return 10
        
        try:
            url = "https://api.steampowered.com/IPlayerService/GetSteamLevel/v1/"
            params = {'key': self.api_key, 'steamid': steam_id}
            
            response = requests.get(url, params=params, timeout=10)
            data = response.json()
            
            if 'response' in data and 'player_level' in data['response']:
                return data['response']['player_level']
        except Exception as e:
            print(f"Не удалось получить уровень для {steam_id}: {str(e)}")
        
        return 10
    
    def get_owned_games(self, steam_id: str) -> dict:
        """Получает список игр и их количество"""
        if not self.api_key:
            return {'game_count': 50, 'games': []}
        
        try:
            url = "https://api.steampowered.com/IPlayerService/GetOwnedGames/v1/"
            params = {
                'key': self.api_key, 
                'steamid': steam_id, 
                'include_appinfo': 1,
                'include_played_free_games': 1
            }
            
            response = requests.get(url, params=params, timeout=15)
            data = response.json()
            
            if 'response' in data:
                return data['response']
        except Exception as e:
            print(f"Не удалось получить игры для {steam_id}: {str(e)}")
        
        return {'game_count': 0, 'games': []}
    
    def get_games_count(self, steam_id: str) -> int:
        """Получает количество игр"""
        games_data = self.get_owned_games(steam_id)
        return games_data.get('game_count', 0)
    
    def get_library_value(self, steam_id: str) -> float:
        """Рассчитывает примерную стоимость библиотеки"""
        if not self.api_key:
            return 500.0
        
        try:
            games_data = self.get_owned_games(steam_id)
            
            if not games_data or 'games' not in games_data:
                return 0
            
            games = games_data['games']
            if not games:
                return 0
            
            return len(games) * 10.0
            
        except Exception as e:
            print(f"Не удалось рассчитать стоимость библиотеки: {str(e)}")
        
        return 0
    
    def get_inventory_value(self, steam_id: str) -> float:
        """Оценивает стоимость инвентаря"""
        if not self.api_key:
            return 100.0
        
        try:
            url = f"https://steamcommunity.com/inventory/{steam_id}/730/2"
            params = {'l': 'russian', 'count': 50}
            
            response = requests.get(url, params=params, timeout=10)
            
            if response.status_code == 200:
                data = response.json()
                if 'assets' in data:
                    item_count = len(data['assets'])
                    return item_count * 5.0
                    
        except Exception as e:
            print(f"Не удалось получить инвентарь: {str(e)}")
        
        return 0
    
    def parse_account(self, account_input: str) -> dict:
        """Основная функция парсинга аккаунта"""
        print(f"\n🔍 Парсинг аккаунта: {account_input[:50]}...")
        
        result = {
            'input': account_input,
            'success': False,
            'error': None,
            'data': {}
        }
        
        try:
            steam_id = self.extract_steam_id(account_input)
            print(f"   SteamID: {steam_id}")
            
            if not steam_id:
                result['error'] = "Не удалось извлечь SteamID"
                return result
            
            player_info = self.get_player_info(steam_id)
            if not player_info:
                result['error'] = "Не удалось получить данные аккаунта"
                return result
            
            steam_level = self.get_steam_level(steam_id)
            games_count = self.get_games_count(steam_id)
            library_value = self.get_library_value(steam_id)
            inventory_value = self.get_inventory_value(steam_id)
            
            result['data'] = {
                'steam_id': steam_id,
                'nickname': player_info.get('personaname', 'Неизвестно'),
                'country': player_info.get('loccountrycode', 'Неизвестно'),
                'avatar': player_info.get('avatarfull', ''),
                'steam_level': steam_level,
                'games_count': games_count,
                'library_value': round(library_value, 2),
                'inventory_value': round(inventory_value, 2),
                'profile_url': player_info.get('profileurl', f'https://steamcommunity.com/profiles/{steam_id}'),
                'parsed_at': datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            }
            
            result['success'] = True
            print(f"   ✅ Успешно: {result['data']['nickname']}")
            
        except Exception as e:
            result['error'] = str(e)
            print(f"   ❌ Ошибка: {str(e)}")
        
        return result
    
    def parse_all_accounts(self):
        """Парсит все аккаунты и сохраняет в БД"""
        print("\n🚀 Начало парсинга всех аккаунтов")
        
        # Создаем сессию парсинга
        session_id = self.db.create_parse_session()
        
        if not session_id:
            print("❌ Не удалось создать сессию парсинга")
            return None, None, []
        
        successful_profiles = []
        failed_profiles = []
        
        for i, account in enumerate(STEAM_ACCOUNTS):
            print(f"\n📊 Аккаунт {i+1}/{len(STEAM_ACCOUNTS)}")
            result = self.parse_account(account)
            
            if result['success']:
                # Получаем или создаем профиль
                profile_id = self.db.get_or_create_profile(
                    result['data']['steam_id'], 
                    result['data']
                )
                
                if profile_id:
                    # Сохраняем снимок данных
                    self.db.save_profile_snapshot(
                        session_id, 
                        profile_id, 
                        result['data']
                    )
                    successful_profiles.append(result['data'])
                    print(f"   ✅ {result['data']['nickname']}")
                else:
                    failed_profiles.append({
                        'account': account,
                        'error': 'Не удалось сохранить профиль в БД'
                    })
                    print(f"   ❌ Ошибка сохранения в БД")
            else:
                failed_profiles.append({
                    'account': account,
                    'error': result.get('error', 'Неизвестная ошибка')
                })
                print(f"   ❌ {result.get('error', 'Ошибка')}")
            
            time.sleep(1)  # Задержка между запросами
        
        # Обновляем статистику сессии
        self.db.update_session_stats(
            session_id,
            len(STEAM_ACCOUNTS),
            len(successful_profiles),
            len(failed_profiles)
        )
        
        # Создаем Word отчет
        if successful_profiles:
            self._create_word_report(session_id, successful_profiles)
        
        return session_id, successful_profiles, failed_profiles
    
    def _create_word_report(self, session_id, profiles_data):
        """Создает Word документ с отчетом"""
        doc = Document()
        
        # Получаем данные сессии
        session = self.db.get_session_by_id(session_id)
        
        # Заголовок
        title = doc.add_heading('Steam Accounts Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Дата и время
        date_para = doc.add_paragraph(f'Дата парсинга: {session["parse_time_display"]}')
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_para.runs[0].bold = True
        
        doc.add_paragraph()
        
        # Сводная статистика
        summary_heading = doc.add_heading('Сводная статистика', 1)
        
        total_profiles = len(profiles_data)
        total_games = sum(p.get('games_count', 0) for p in profiles_data)
        total_level = sum(p.get('steam_level', 0) for p in profiles_data)
        avg_level = total_level / total_profiles if total_profiles > 0 else 0
        total_value = sum(p.get('library_value', 0) + p.get('inventory_value', 0) for p in profiles_data)
        
        summary_table = doc.add_table(rows=5, cols=2)
        summary_table.style = 'Light Grid Accent 1'
        summary_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        data = [
            ("Количество аккаунтов", str(total_profiles)),
            ("Всего игр в библиотеках", str(total_games)),
            ("Средний уровень Steam", f"{avg_level:.1f}"),
            ("Суммарный уровень", str(total_level)),
            ("Общая стоимость", f"${total_value:,.2f}")
        ]
        
        for i, (label, value) in enumerate(data):
            summary_table.cell(i, 0).text = label
            summary_table.cell(i, 1).text = value
            summary_table.cell(i, 0).paragraphs[0].runs[0].bold = True
        
        doc.add_paragraph()
        
        # Детальная информация по каждому аккаунту
        details_heading = doc.add_heading('Детальная информация по аккаунтам', 1)
        
        for i, profile in enumerate(profiles_data, 1):
            account_heading = doc.add_heading(f'Аккаунт {i}: {profile.get("nickname", "Неизвестно")}', 2)
            
            # Основная информация
            info_para = doc.add_paragraph()
            info_para.add_run(f"SteamID: ").bold = True
            info_para.add_run(f'{profile.get("steam_id", "N/A")}\n')
            
            info_para.add_run(f"Страна: ").bold = True
            info_para.add_run(f'{profile.get("country", "Неизвестно")}\n')
            
            info_para.add_run(f"Уровень Steam: ").bold = True
            info_para.add_run(f'{profile.get("steam_level", 0)}\n')
            
            info_para.add_run(f"Дата парсинга: ").bold = True
            info_para.add_run(f'{profile.get("parsed_at", "N/A")}\n')
            
            # Статистика в таблице
            stats_table = doc.add_table(rows=3, cols=2)
            stats_table.style = 'Light Grid Accent 2'
            
            stats_data = [
                ("Игр в библиотеке", str(profile.get('games_count', 0))),
                ("Стоимость библиотеки", f"${profile.get('library_value', 0):,.2f}"),
                ("Стоимость инвентаря", f"${profile.get('inventory_value', 0):,.2f}")
            ]
            
            for row, (label, value) in enumerate(stats_data):
                stats_table.cell(row, 0).text = label
                stats_table.cell(row, 1).text = value
                stats_table.cell(row, 0).paragraphs[0].runs[0].bold = True
            
            doc.add_paragraph()
            
            # Ссылка на профиль
            link_para = doc.add_paragraph()
            link_para.add_run("Ссылка на профиль: ").bold = True
            link_para.add_run(f'{profile.get("profile_url", "")}')
            
            if i < len(profiles_data):
                doc.add_paragraph("—" * 50)
        
        # Сохраняем документ
        filename = f"Steam_Report_{session['timestamp_str']}.docx"
        filepath = WORD_REPORTS_DIR / filename
        doc.save(filepath)
        print(f"✅ Word отчет сохранен: {filepath}")

# ==================== ФУНКЦИИ ДЛЯ STREAMLIT ====================

def format_currency(value):
    """Форматирует валюту"""
    return f"${float(value):,.2f}"

def main():
    st.set_page_config(
        page_title="Steam Parser with PostgreSQL",
        page_icon="🎮",
        layout="wide"
    )
    
    st.title("🎮 Steam Account Parser with PostgreSQL")
    
    # Инициализация сервисов
    db = DatabaseManager(DB_CONFIG)
    parser = SteamParser(db)
    
    # Боковая панель
    with st.sidebar:
        st.header("⚙️ Управление")
        
        # Ручной парсинг
        if st.button("🚀 Запустить парсинг сейчас", type="primary", use_container_width=True):
            with st.spinner("Парсинг аккаунтов..."):
                session_id, successful, failed = parser.parse_all_accounts()
                
                if successful:
                    st.success(f"✅ Успешно обработано: {len(successful)} аккаунтов")
                    if failed:
                        st.warning(f"⚠️ Ошибок: {len(failed)}")
                    
                    # Показываем детали
                    with st.expander("📋 Детали парсинга"):
                        for profile in successful:
                            st.write(f"✅ {profile['nickname']} (Уровень: {profile['steam_level']})")
                        for fail in failed:
                            st.write(f"❌ {fail['account'][:50]}...: {fail['error']}")
                    
                    st.info(f"ID сессии: {session_id}")
                else:
                    st.error("❌ Не удалось обработать ни одного аккаунта")
        
        st.divider()
        
        # История сессий
        st.header("📅 История парсинга")
        
        sessions = db.get_sessions(limit=50)
        
        if sessions:
            st.success(f"📊 Всего сессий: {len(sessions)}")
            
            # Создаем словарь для выбора
            session_options = {}
            for s in sessions:
                label = f"{s['parse_time_display']} | ✅ {s['successful_profiles']}/{s['total_profiles']} | 💰 ${float(s['grand_total_value']):,.0f}"
                session_options[label] = s['session_id']
            
            selected_label = st.selectbox(
                "Выберите сессию для просмотра:",
                list(session_options.keys())
            )
            
            if st.button("📖 Показать выбранную сессию", use_container_width=True):
                st.session_state.selected_session_id = session_options[selected_label]
                st.rerun()
            
            # Кнопка для просмотра всех сессий
            if st.button("📊 Показать все сессии", use_container_width=True):
                st.session_state.show_all_sessions = True
                st.rerun()
        else:
            st.info("📭 Сессий пока нет")
        
        st.divider()
        
        # Общая статистика БД
        stats = db.get_stats()
        if stats:
            st.header("💾 Статистика БД")
            st.metric("Всего профилей", stats['total_profiles'])
            st.metric("Всего сессий", stats['total_sessions'])
            st.metric("Всего снимков", stats['total_snapshots'])
            if stats['total_value']:
                st.metric("Общая стоимость", format_currency(stats['total_value']))
    
    # Основная область
    if 'selected_session_id' in st.session_state:
        # Показываем выбранную сессию
        session_id = st.session_state.selected_session_id
        session_data = db.get_session_by_id(session_id)
        
        if session_data:
            st.header(f"📄 Сессия от {session_data['parse_time_display']}")
            
            # Сводная статистика
            col1, col2, col3, col4, col5 = st.columns(5)
            with col1:
                st.metric("Всего аккаунтов", session_data['total_profiles'])
            with col2:
                st.metric("Успешно", session_data['successful_profiles'])
            with col3:
                st.metric("Всего игр", session_data['total_games'] or 0)
            with col4:
                st.metric("Средний уровень", f"{float(session_data['avg_level']):.1f}")
            with col5:
                st.metric("Общая стоимость", format_currency(session_data['grand_total_value']))
            
            # Получаем профили сессии
            profiles = db.get_session_profiles(session_id)
            
            if profiles:
                # Графики
                st.subheader("📊 Визуализация данных")
                
                tab1, tab2, tab3 = st.tabs(["📈 Стоимость", "🎮 Игры", "🌍 Страны"])
                
                with tab1:
                    # График стоимости
                    fig = go.Figure()
                    
                    names = [p['nickname'] for p in profiles]
                    library_values = [float(p['library_value']) for p in profiles]
                    inventory_values = [float(p['inventory_value']) for p in profiles]
                    
                    fig.add_trace(go.Bar(
                        name='Библиотека',
                        x=names,
                        y=library_values,
                        marker_color='rgb(55, 83, 109)'
                    ))
                    
                    fig.add_trace(go.Bar(
                        name='Инвентарь',
                        x=names,
                        y=inventory_values,
                        marker_color='rgb(26, 118, 255)'
                    ))
                    
                    fig.update_layout(
                        title="Стоимость библиотеки и инвентаря по аккаунтам",
                        xaxis_title="Аккаунт",
                        yaxis_title="Стоимость ($)",
                        barmode='group'
                    )
                    
                    st.plotly_chart(fig, use_container_width=True)
                
                with tab2:
                    # График количества игр
                    fig = px.bar(
                        x=[p['nickname'] for p in profiles],
                        y=[p['games_count'] for p in profiles],
                        title="Количество игр в библиотеке",
                        labels={'x': 'Аккаунт', 'y': 'Количество игр'}
                    )
                    st.plotly_chart(fig, use_container_width=True)
                
                with tab3:
                    # Статистика по странам
                    country_counts = {}
                    for p in profiles:
                        country = p['country'] or 'Неизвестно'
                        country_counts[country] = country_counts.get(country, 0) + 1
                    
                    fig = px.pie(
                        values=list(country_counts.values()),
                        names=list(country_counts.keys()),
                        title="Распределение по странам"
                    )
                    st.plotly_chart(fig, use_container_width=True)
                
                # Детальная информация
                st.subheader("👤 Детальная информация по аккаунтам")
                
                for profile in profiles:
                    with st.expander(f"🎮 {profile['nickname']}"):
                        col1, col2 = st.columns([1, 3])
                        
                        with col1:
                            if profile['avatar_url']:
                                st.image(profile['avatar_url'], width=100)
                            st.metric("Уровень", profile['snapshot_level'])
                        
                        with col2:
                            st.write(f"**Страна:** {profile['country'] or 'Неизвестно'}")
                            st.write(f"**SteamID:** `{profile['steam_id']}`")
                            st.write(f"**Игр в библиотеке:** {profile['games_count']}")
                            st.write(f"**Стоимость библиотеки:** {format_currency(profile['library_value'])}")
                            st.write(f"**Стоимость инвентаря:** {format_currency(profile['inventory_value'])}")
                            st.write(f"**Общая стоимость:** {format_currency(profile['total_value'])}")
                            st.write(f"**Ссылка:** {profile['profile_url']}")
                            
                            # Кнопка для просмотра истории профиля
                            if st.button(f"📈 История профиля", key=f"history_{profile['id']}"):
                                st.session_state.selected_profile_id = profile['id']
                                st.session_state.selected_profile_name = profile['nickname']
                                st.rerun()
                
                # Кнопка возврата
                if st.button("⬅️ Назад к списку сессий"):
                    del st.session_state.selected_session_id
                    if 'selected_profile_id' in st.session_state:
                        del st.session_state.selected_profile_id
                    st.rerun()
        
        # Показываем историю профиля если выбрана
        if 'selected_profile_id' in st.session_state:
            st.divider()
            st.subheader(f"📈 История профиля: {st.session_state.selected_profile_name}")
            
            history = db.get_profile_history(st.session_state.selected_profile_id)
            
            if history:
                # Создаем DataFrame для графика
                df = pd.DataFrame(history)
                df['parse_time'] = pd.to_datetime(df['parse_time'])
                
                # График изменения уровня
                fig = px.line(
                    df, 
                    x='parse_time', 
                    y='steam_level',
                    title="Изменение уровня Steam",
                    labels={'parse_time': 'Дата', 'steam_level': 'Уровень'}
                )
                st.plotly_chart(fig, use_container_width=True)
                
                # График изменения стоимости
                fig = px.line(
                    df, 
                    x='parse_time', 
                    y=['library_value', 'inventory_value', 'total_value'],
                    title="Изменение стоимости",
                    labels={'parse_time': 'Дата', 'value': 'Стоимость ($)'}
                )
                st.plotly_chart(fig, use_container_width=True)
                
                # Таблица истории
                st.dataframe(
                    df[['parse_time_display', 'steam_level', 'games_count', 
                        'library_value', 'inventory_value', 'total_value']],
                    use_container_width=True
                )
            else:
                st.info("Нет данных истории для этого профиля")
    
    elif 'show_all_sessions' in st.session_state:
        # Показываем все сессии в виде таблицы
        st.header("📊 Все сессии парсинга")
        
        sessions = db.get_sessions(limit=100)
        
        if sessions:
            df = pd.DataFrame(sessions)
            df['parse_time'] = pd.to_datetime(df['parse_time'])
            
            # Форматируем для отображения
            display_df = df[['parse_time_display', 'total_profiles', 'successful_profiles',
                           'failed_profiles', 'total_games', 'avg_level', 'grand_total_value']].copy()
            
            display_df.columns = ['Время парсинга', 'Всего', 'Успешно', 'Ошибок',
                                'Всего игр', 'Ср. уровень', 'Общая стоимость']
            
            display_df['Общая стоимость'] = display_df['Общая стоимость'].apply(format_currency)
            display_df['Ср. уровень'] = display_df['Ср. уровень'].apply(lambda x: f"{float(x):.1f}")
            
            st.dataframe(display_df, use_container_width=True)
            
            # График по сессиям
            fig = px.line(
                df, 
                x='parse_time', 
                y='grand_total_value',
                title="Динамика общей стоимости",
                labels={'parse_time': 'Дата', 'grand_total_value': 'Общая стоимость ($)'}
            )
            st.plotly_chart(fig, use_container_width=True)
            
            if st.button("⬅️ Назад"):
                del st.session_state.show_all_sessions
                st.rerun()
        else:
            st.info("Нет данных")
    
    else:
        # Главная страница
        st.header("📊 Система мониторинга Steam аккаунтов с PostgreSQL")
        
        col1, col2 = st.columns(2)
        
        with col1:
            st.info("""
            ### 🎯 Возможности системы:
            1. **PostgreSQL хранилище** всех данных
            2. **История изменений** каждого профиля
            3. **Сравнение сессий** во времени
            4. **Графики и аналитика** в реальном времени
            5. **Экспорт в Word** и JSON
            """)
        
        with col2:
            # Последние сессии
            st.subheader("🕒 Последние сессии")
            sessions = db.get_sessions(limit=5)
            
            if sessions:
                for s in sessions:
                    st.write(f"📅 {s['parse_time_display']}")
                    st.write(f"   ✅ {s['successful_profiles']}/{s['total_profiles']} | 💰 {format_currency(s['grand_total_value'])}")
                    st.divider()
            else:
                st.write("Нет данных. Запустите парсинг!")
        
        # Статус подключения
        st.divider()
        st.subheader("🔧 Статус системы")
        
        col1, col2 = st.columns(2)
        
        with col1:
            if STEAM_API_KEY:
                st.success("✅ Steam API: OK")
            else:
                st.error("❌ Steam API: Нет ключа")
        
        with col2:
            try:
                stats = db.get_stats()
                st.success(f"✅ PostgreSQL: OK (профилей: {stats['total_profiles']})")
            except:
                st.error("❌ PostgreSQL: Ошибка подключения")

# ==================== ФУНКЦИЯ ДЛЯ АВТОМАТИЧЕСКОГО ПАРСИНГА ====================

def run_auto_parse():
    """Функция для запуска из cron"""
    print("=" * 50)
    print("Запуск авто-парсинга Steam аккаунтов")
    print(f"Время запуска: {datetime.now().strftime('%d.%m.%Y %H:%M:%S')}")
    print("=" * 50)
    
    db = DatabaseManager(DB_CONFIG)
    parser = SteamParser(db)
    
    session_id, successful, failed = parser.parse_all_accounts()
    
    if successful:
        print(f"✅ Отчет сохранен в БД (session_id: {session_id})")
        print(f"   Обработано: {len(successful)}/{len(STEAM_ACCOUNTS)}")
    else:
        print("❌ Не удалось обработать ни одного аккаунта")
    
    print("=" * 50)

# ==================== ЗАПУСК ====================

if __name__ == "__main__":
    import sys
    if len(sys.argv) > 1 and sys.argv[1] == "--auto":
        run_auto_parse()
    else:
        main()